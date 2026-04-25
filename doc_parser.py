"""
文档内容提取工具 v2
改进：按段落/章节分块，支持重新索引
"""
import io
import re
import sqlite3
from config import DATABASE_URL


def extract_text_from_pdf(file_bytes: bytes) -> str:
    try:
        import PyPDF2
        reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
        texts = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                texts.append(text)
        return "\n".join(texts)
    except:
        return ""


def extract_text_from_docx(file_bytes: bytes) -> str:
    try:
        import docx
        doc = docx.Document(io.BytesIO(file_bytes))
        texts = []
        for para in doc.paragraphs:
            if para.text.strip():
                texts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    texts.append(row_text)
        return "\n".join(texts)
    except:
        return ""


def extract_text_from_excel(file_bytes: bytes) -> str:
    try:
        import openpyxl
        wb = openpyxl.load_workbook(io.BytesIO(file_bytes), read_only=True, data_only=True)
        texts = []
        for sheet in wb.worksheets:
            texts.append(f"[Sheet: {sheet.title}]")
            for row in sheet.iter_rows(values_only=True):
                row_vals = [str(c) for c in row if c is not None]
                if row_vals:
                    texts.append(" | ".join(row_vals))
        wb.close()
        return "\n".join(texts)
    except:
        return ""


def extract_text_from_bytes(file_bytes: bytes, file_type: str) -> str:
    ft = file_type.lower()
    if ft == ".pdf":
        return extract_text_from_pdf(file_bytes)
    elif ft in (".docx", ".doc"):
        return extract_text_from_docx(file_bytes)
    elif ft in (".xlsx", ".xls"):
        return extract_text_from_excel(file_bytes)
    elif ft in (".txt", ".md", ".csv", ".xml", ".json"):
        try:
            return file_bytes.decode("utf-8", errors="ignore")
        except:
            return file_bytes.decode("gbk", errors="ignore")
    return ""


# ─── 智能分块 ─────────────────────────────────────

def smart_chunk(text: str, max_len: int = 800, overlap: int = 100) -> list:
    """
    智能分块：按段落/章节切分，保留上下文重叠
    比简单切分效果好很多
    """
    if not text.strip():
        return []

    # 先按空行/标题切段
    sections = re.split(r'\n{2,}|\n(?=[一二三四五六七八九十]+[、.])|\n(?=\d+[、.])|\n(?=[A-Z][\.\s])', text)
    sections = [s.strip() for s in sections if s.strip()]

    chunks = []
    current = []
    current_len = 0

    for section in sections:
        # 如果单个段落超长，按句子切
        if len(section) > max_len:
            if current:
                chunks.append("\n".join(current))
                current = []
                current_len = 0
            sentences = re.split(r'(?<=[。！？；\n])', section)
            sub_chunk = []
            sub_len = 0
            for sent in sentences:
                if sub_len + len(sent) > max_len and sub_chunk:
                    chunks.append("".join(sub_chunk))
                    # 保留最后 overlap 字符作为上下文
                    overlap_text = "".join(sub_chunk)[-overlap:]
                    sub_chunk = [overlap_text, sent]
                    sub_len = len(overlap_text) + len(sent)
                else:
                    sub_chunk.append(sent)
                    sub_len += len(sent)
            if sub_chunk:
                chunks.append("".join(sub_chunk))
            continue

        if current_len + len(section) > max_len and current:
            chunks.append("\n".join(current))
            # 保留最后一个段落作为重叠
            if current:
                overlap_text = current[-1]
                current = [overlap_text] if len(overlap_text) < overlap else []
                current_len = sum(len(c) for c in current)
            else:
                current = []
                current_len = 0

        current.append(section)
        current_len += len(section)

    if current:
        chunks.append("\n".join(current))

    return chunks if chunks else [text[:max_len]]


# ─── 全文检索索引 ─────────────────────────────────

def init_fts():
    db_path = DATABASE_URL.replace("sqlite:///", "")
    conn = sqlite3.connect(db_path)
    conn.execute("""
        CREATE VIRTUAL TABLE IF NOT EXISTS doc_fts USING fts5(
            doc_id, title, content, product_model, doc_number,
            tokenize='unicode61'
        )
    """)
    conn.commit()
    conn.close()


def index_document(doc_id: int, title: str, content: str,
                    product_model: str = "", doc_number: str = ""):
    if not content.strip():
        return
    db_path = DATABASE_URL.replace("sqlite:///", "")
    conn = sqlite3.connect(db_path)
    conn.execute("DELETE FROM doc_fts WHERE doc_id = ?", (str(doc_id),))
    chunks = smart_chunk(content, max_len=800, overlap=100)
    for chunk in chunks:
        conn.execute(
            "INSERT INTO doc_fts (doc_id, title, content, product_model, doc_number) VALUES (?, ?, ?, ?, ?)",
            (str(doc_id), title, chunk, product_model, doc_number)
        )
    conn.commit()
    conn.close()


def reindex_all_documents():
    """重新索引所有文档（用于升级后重建索引）"""
    import sys
    sys.path.insert(0, '.')
    from encryption import decrypt_file

    db_path = DATABASE_URL.replace("sqlite:///", "")
    conn = sqlite3.connect(db_path)

    # 清空旧索引
    conn.execute("DELETE FROM doc_fts")

    # 获取所有文档
    conn2 = sqlite3.connect(db_path)
    conn2.row_factory = sqlite3.Row
    cursor = conn2.execute("SELECT id, title, filename, stored_name, file_type, product_model, doc_number FROM documents")
    docs = cursor.fetchall()
    conn2.close()

    import os
    from pathlib import Path
    from config import DATA_DIR

    count = 0
    for doc in docs:
        stored = DATA_DIR / doc['stored_name']
        if not stored.exists():
            continue
        try:
            plaintext = decrypt_file(doc['stored_name'])
            text = extract_text_from_bytes(plaintext, doc['file_type'] or '')
            if text.strip():
                chunks = smart_chunk(text)
                for chunk in chunks:
                    conn.execute(
                        "INSERT INTO doc_fts (doc_id, title, content, product_model, doc_number) VALUES (?, ?, ?, ?, ?)",
                        (str(doc['id']), doc['title'], chunk, doc['product_model'] or '', doc['doc_number'] or '')
                    )
                count += 1
        except:
            pass

    conn.commit()
    conn.close()
    return count


def search_documents(query: str, limit: int = 10) -> list:
    db_path = DATABASE_URL.replace("sqlite:///", "")
    conn = sqlite3.connect(db_path)
    try:
        # 使用 BM25 排序 + snippet 高亮
        cursor = conn.execute(
            """SELECT doc_id, title, snippet(doc_fts, 2, '【', '】', '...', 80) as snip,
                      rank, product_model, doc_number
               FROM doc_fts WHERE doc_fts MATCH ? ORDER BY rank LIMIT ?""",
            (query, limit)
        )
        results = []
        seen = set()
        for row in cursor:
            did = int(row[0])
            if did in seen:
                continue
            seen.add(did)
            results.append({
                "doc_id": did,
                "title": row[1],
                "snippet": row[2],
                "rank": row[3],
                "product_model": row[4],
                "doc_number": row[5],
            })
        return results
    except:
        return []
    finally:
        conn.close()


def search_document_context(query: str, doc_id: int = None, limit: int = 5) -> str:
    """
    搜索文档内容：先尝试 FTS5，失败则用 LIKE 模糊搜索
    """
    db_path = DATABASE_URL.replace("sqlite:///", "")
    conn = sqlite3.connect(db_path)
    try:
        contexts = []
        total_len = 0

        # 先尝试 FTS5 搜索
        try:
            if doc_id:
                cursor = conn.execute(
                    """SELECT doc_id, title, content FROM doc_fts
                       WHERE doc_fts MATCH ? AND doc_id = ? LIMIT ?""",
                    (query, str(doc_id), limit)
                )
            else:
                cursor = conn.execute(
                    """SELECT doc_id, title, content FROM doc_fts
                       WHERE doc_fts MATCH ? ORDER BY rank LIMIT ?""",
                    (query, limit)
                )
            for row in cursor:
                content = row[2][:1200]
                contexts.append(f"[文档: {row[1]}]\n{content}")
                total_len += len(content)
                if total_len > 6000:
                    break
        except:
            pass

        # FTS 没结果，用 LIKE 搜索（中文友好）
        if not contexts:
            keywords = [kw.strip() for kw in query.replace("，", " ").replace(",", " ").split() if kw.strip()]
            if not keywords:
                keywords = [query]

            for kw in keywords[:3]:
                if doc_id:
                    cursor = conn.execute(
                        """SELECT doc_id, title, content FROM doc_fts
                           WHERE (content LIKE ? OR title LIKE ?) AND doc_id = ? LIMIT ?""",
                        (f"%{kw}%", f"%{kw}%", str(doc_id), limit)
                    )
                else:
                    cursor = conn.execute(
                        """SELECT doc_id, title, content FROM doc_fts
                           WHERE content LIKE ? OR title LIKE ? LIMIT ?""",
                        (f"%{kw}%", f"%{kw}%", limit)
                    )
                for row in cursor:
                    content = row[2][:1200]
                    entry = f"[文档: {row[1]}]\n{content}"
                    if entry not in contexts:
                        contexts.append(entry)
                        total_len += len(content)
                    if total_len > 6000:
                        break
                if total_len > 6000:
                    break

        # 仍然没结果，返回文档目录摘要
        if not contexts:
            try:
                if doc_id:
                    cursor = conn.execute(
                        "SELECT DISTINCT doc_id, title FROM doc_fts WHERE doc_id = ?",
                        (str(doc_id),)
                    )
                else:
                    cursor = conn.execute(
                        "SELECT DISTINCT doc_id, title FROM doc_fts LIMIT 10"
                    )
                doc_list = []
                for row in cursor:
                    doc_list.append(f"  - {row[1]} (ID: {row[0]})")
                if doc_list:
                    contexts.append(
                        "[当前已上传的文档]\n" + "\n".join(doc_list) +
                        "\n\n提示：未在文档中找到与问题直接匹配的内容，请尝试更具体的关键词，"
                        "或选择指定文档后重新提问。"
                    )
            except:
                pass

        return "\n\n---\n\n".join(contexts)
    except:
        return ""
    finally:
        conn.close()


def get_all_doc_content(doc_id: int) -> str:
    db_path = DATABASE_URL.replace("sqlite:///", "")
    conn = sqlite3.connect(db_path)
    try:
        cursor = conn.execute(
            "SELECT content FROM doc_fts WHERE doc_id = ?", (str(doc_id),)
        )
        texts = [row[0] for row in cursor]
        return "\n".join(texts)
    except:
        return ""
    finally:
        conn.close()
